from pathlib import Path
from docx import Document  # package name: python-docx
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
import re
from botok import TokChunks, SylComponents
from .composition_full import roots, exceptions, ends


def get_composition(part):
    if part in roots:
        return roots[part]
    if part in exceptions:
        return exceptions[part]
    if part in ends:
        return ends[part]
    return ''


def get_syllables(string):
    c = TokChunks(string)
    c.serve_syls_to_trie()
    total = []
    for s, info in c.chunks:
        _, start, length = info
        if s:
            s_comp = SylComponents()
            syl = ''.join([string[x] for x in s])
            parts = s_comp.get_parts(syl)
            comp = [get_composition(p) for p in parts]
            if len(s) < len(string[start:start+length]):
                comp.append('X')
                syl += '་'
            total.append((syl, ''.join(comp)))
        else:
            total.append((string[start:start+length], length * 'X'))
    return total


def create_docx(chunks, path):
    # CONFIG
    indentation = 0.5
    tibetan_style = {
        'font': 'Jomolhari',
        'size': 11
    }

    word1 = {
        # 1. mingzhi
        'm': {'color': (0, 153, 0)},
        # 2. vowel
        'v': {'color': (0, 204, 0)},
        # 3. superscript and subscript
        't': {'color': (0, 255, 0)},
        'b': {'color': (0, 255, 0)},
        # 4. prefix and suffix
        'p': {'color': (102, 255, 102)},
        's': {'color': (102, 255, 102)},
        # 5. second suffix and others
        'S': {'color': (192, 192, 192)},
        'X': {'color': (192, 192, 192)}
    }
    word2 = {
        # 1.
        'm': {'color': (76, 0, 153)},
        # 2.
        'v': {'color': (102, 0, 204)},
        # 3.
        't': {'color': (127, 0, 255)},
        'b': {'color': (127, 0, 255)},
        # 4.
        'p': {'color': (178, 102, 255)},
        's': {'color': (178, 102, 255)},
        # 5.
        'S': {'color': (192, 192, 192)},
        'X': {'color': (192, 192, 192)}
    }

    pedurma_style = {
        'color': (112, 128, 144),
        'font': 'Jomolhari',
        'size': 8
    }
    semantic_style = {
        'font': 'Free Mono',
        'size': 10
    }
    communicative_style = {
        'font': 'Gentium',
        'size': 12
    }

    document = Document()
    styles = document.styles

    # TIBETAN
    bo_style = styles.add_style('Tibetan', WD_STYLE_TYPE.CHARACTER)
    bo_font = bo_style.font
    bo_font.name = tibetan_style['font']
    bo_font.size = Pt(tibetan_style['size'])
    # PEYDURMA NOTES
    prefix_style = styles.add_style('Peydurma Notes', WD_STYLE_TYPE.CHARACTER)
    prefix_font = prefix_style.font
    c = pedurma_style['color']
    prefix_font.color.rgb = RGBColor(c[0], c[1], c[2])

    # COMMUNICATIVE VERSION
    com_style = styles.add_style('Communicative', WD_STYLE_TYPE.CHARACTER)
    com_font = com_style.font
    com_font.name = communicative_style['font']
    com_font.size = Pt(communicative_style['size'])

    # SEMANTIC VERSION
    sem_style = styles.add_style('Semantic', WD_STYLE_TYPE.CHARACTER)
    sem_style.base_style = styles['Normal']
    sem_font = sem_style.font
    sem_font.name = semantic_style['font']
    sem_font.size = Pt(semantic_style['size'])

    # COMMUNICATIVE PARAGRAPH
    com_par_style = styles.add_style('Com. paragraph', WD_STYLE_TYPE.PARAGRAPH)
    com_par_style.paragraph_format.space_before = Cm(0)
    com_par_style.paragraph_format.space_after = Cm(0)

    # OTHER PARAGRAPH
    other_par_style = styles.add_style('Other paragraph', WD_STYLE_TYPE.PARAGRAPH)
    other_par_style.paragraph_format.space_before = Cm(0)
    other_par_style.paragraph_format.space_after = Cm(1)
    other_par_style.paragraph_format.left_indent = Cm(indentation)
    other_par_style.paragraph_format.line_spacing = WD_LINE_SPACING.SINGLE

    com_p = document.add_paragraph()

    for chunk in chunks:
        chars, codes = chunk
        for i in range(len(chars)):
            char, code = chars[i], codes[i]
            run = com_p.add_run(char)
            c = word1[code]['color']
            run.font.color.rgb = RGBColor(c[0], c[1], c[2])

    out_path = path.parent / (path.stem + '.docx')
    document.save(str(out_path))

truc = 'བསྒྲུབས་བསྒྲུབས། །'
output = Path('./truc')
syls = get_syllables(truc)
create_docx(syls, output)