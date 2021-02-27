from pathlib import Path
from docx import Document  # package name: python-docx
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
import re
from botok import TokChunks, SylComponents
from composition import roots, exceptions, ends

fixed = {
    'prefix': (220, 220, 220),
    'suffix2': (220, 220, 220),
    'other': (0, 0, 128)
}

color_profiles = {
    # syllable structure:
    #   p: prefix
    #   m: main stack
    #   s: suffix
    #   S: 2nd suffix
    #   X: punct etc.
    'strong': {
        'word1': {  # purple, light purple
            'p': {'color': fixed['prefix']},
            'm': {'color': (138, 43, 226)},
            's': {'color': (162, 162, 208)},
            'S': {'color': fixed['suffix2']},
            'X': {'color': fixed['other']}},

        'word2': {  # green, light green
            'p': {'color': fixed['prefix']},
            'm': {'color': (3, 192, 60)},
            's': {'color': (172, 225, 175)},
            'S': {'color': fixed['suffix2']},
            'X': {'color': fixed['other']}},

        'other': {  # dark brown everywhere
            'm': {'color': (149, 74, 41)},
            'X': {'color': fixed['other']}}
    },

    'dark': {
        'word1': {  # purple, light purple
            'p': {'color': fixed['prefix']},
            'm': {'color': (59, 1, 112)},
            's': {'color': (145, 124, 163)},
            'S': {'color': fixed['suffix2']},
            'X': {'color': fixed['other']}},

        'word2': {  # green, light green
            'p': {'color': fixed['prefix']},
            'm': {'color': (2, 82, 26)},
            's': {'color': (81, 168, 108)},
            'S': {'color': fixed['suffix2']},
            'X': {'color': fixed['other']}},

        'other': {  # dark brown everywhere
            'm': {'color': (115, 54, 28)},
            'X': {'color': fixed['other']}}
    }
}


def get_composition(halves):
    begin, end = halves
    parts = []
    if begin in roots:
        parts.extend(roots[begin])
    if begin in exceptions:
        parts.extend(exceptions[begin])
    if end in ends:
        compos = ends[end]
        if compos[0][0] == 'v':
            parts[-1] = (parts[-1][0], parts[-1][1] + compos[0][1])
            compos = compos[1:]
        parts.extend(compos)
    return parts


def make_chunks(string, mode='spaces'):
    string = string.replace('་ ', '་-')
    if mode == 'spaces':
        regexp = '(-)'
    elif mode == 'no_spaces':
        regexp = '-'
    else:
        raise ValueError('either spaces or non_spaces')
    chunks = []
    for c in re.split(regexp, string):
        if c == '-':
            if not chunks:
                chunks.append(' ')
            else:
                chunks[-1] += c
        else:
            chunks.append(c)
    return chunks


def get_syllables(string, color_switch, mode):
    total = []
    for chunk in make_chunks(string, mode=mode):
        if color_switch == 1:
            color_switch = 2
        elif color_switch == 2:
            color_switch = 1
        c = TokChunks(chunk)
        c.serve_syls_to_trie()
        for s, info in c.chunks:
            _, start, length = info
            if s:
                s_comp = SylComponents()
                syl = ''.join([chunk[x] for x in s])
                parts = s_comp.get_parts(syl)
                if parts:
                    comp = get_composition(parts)
                    if len(s) < len(chunk[start:start + length]):
                        punct = ''.join([chunk[i] for i in range(s[-1] + 1, start + length)])
                        comp.append(('X', punct))
                    total.append((color_switch, comp))
                else:
                    comp = [('m', syl)]
                    if len(s) < len(chunk[start:start + length]):
                        punct = ''.join([chunk[i] for i in range(s[-1] + 1, start + length)])
                        comp.append(('X', punct))
                    total.append((3, comp))
                    if color_switch == 1:
                        color_switch = 2
                    elif color_switch == 2:
                        color_switch = 1
            else:
                total.append((color_switch, [('X', chunk[start:start+length])]))


    return total, color_switch


def initiate_docx():
    # CONFIG
    indentation = 0.5
    tibetan_style = {
        'font': 'Monlam Uni OuChan2',
        'size': 28
    }

    document = Document()
    styles = document.styles

    # TIBETAN
    bo_style = styles.add_style('Tibetan', WD_STYLE_TYPE.CHARACTER)
    bo_font = bo_style.font
    bo_font.name = tibetan_style['font']
    bo_font.size = Pt(tibetan_style['size'])
    return document


def create_docx(lines, path, profile='strong', mode='spaces'):
    document = initiate_docx()

    com_p = document.add_paragraph()
    switch = 2
    for line in lines:
        if line == '':
            com_p.add_run().add_break()
        else:
            chunks, switch = get_syllables(line, switch, mode=mode)
            for switch, chunk in chunks:
                for code, part in chunk:
                    run = com_p.add_run(part, style=document.styles['Tibetan'])
                    if switch == 1:
                        c = color_profiles[profile]['word1'][code]['color']
                    elif switch == 2:
                        c = color_profiles[profile]['word2'][code]['color']
                    elif switch == 3:
                        c = color_profiles[profile]['other'][code]['color']
                    else:
                        raise ValueError('can be either 1, 2 or 3')
                    run.font.color.rgb = RGBColor(c[0], c[1], c[2])
            com_p.add_run().add_break()

    out_path = path.parent / (path.stem + '.docx')
    document.save(str(out_path))


if __name__ == '__main__':
    in_path = Path('tib')
    for f in in_path.glob('*.txt'):
        print(f)
        lines = f.read_text(encoding='utf-8').strip().split('\n')
        output = Path('out') / (f.stem + '.docx')
        create_docx(lines, output, profile='dark', mode='no_spaces')
