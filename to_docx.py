from pathlib import Path
from docx import Document  # package name: python-docx
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
import re
from botok import TokChunks, SylComponents
from composition import roots, exceptions, ends


def get_composition(halves):
    begin, end = halves
    parts = []
    if begin in roots:
        parts.extend(roots[begin])
    if begin in exceptions:
        parts.extend(exceptions[begin])
    if end in ends:
        compos = ends[end]
        if compos[0][0] == 'v':
            parts[-1] = (parts[-1][0], parts[-1][1] + compos[0][1])
            compos = compos[1:]
        parts.extend(compos)
    return parts


def make_chunks(string):
    chunks = []
    for c in re.split('( )', string):
        if c == ' ':
            if not chunks:
                chunks.append(c)
            else:
                chunks[-1] += c
        else:
            chunks.append(c)
    return chunks


def get_syllables(string):
    word_marker = 2
    total = []
    for chunk in make_chunks(string):
        if word_marker == 1:
            word_marker = 2
        elif word_marker == 2:
            word_marker = 1
        c = TokChunks(chunk)
        c.serve_syls_to_trie()
        for s, info in c.chunks:
            _, start, length = info
            if s:
                s_comp = SylComponents()
                syl = ''.join([chunk[x] for x in s])
                parts = s_comp.get_parts(syl)
                comp = get_composition(parts)
                if len(s) < len(chunk[start:start+length]):
                    punct = ''.join([chunk[i] for i in range(s[-1] + 1, start + length)])
                    comp.append(('X', punct))
                total.append((word_marker, comp))
            else:
                total.append((word_marker, [('X', chunk[start:start+length])]))
    return total


def initiate_docx():
    # CONFIG
    indentation = 0.5
    tibetan_style = {
        'font': 'Jomolhari',
        'size': 11
    }

    word1 = {  # pmvsS
        # 1. mingzhi stack
        'm': {'color': (0, 153, 0)},
        # 2. suffix
        's': {'color': (102, 255, 102)},
        # 3. prefix, second suffix and others
        'p': {'color': (192, 192, 192)},
        'S': {'color': (192, 192, 192)},
        'X': {'color': (192, 192, 192)}
    }
    word2 = {
        # 1.
        'm': {'color': (76, 0, 153)},
        # 2.
        's': {'color': (178, 102, 255)},
        # 3.
        'p': {'color': (192, 192, 192)},
        'S': {'color': (192, 192, 192)},
        'X': {'color': (192, 192, 192)}
    }

    pedurma_style = {
        'color': (112, 128, 144),
        'font': 'Jomolhari',
        'size': 8
    }
    semantic_style = {
        'font': 'Free Mono',
        'size': 10
    }
    communicative_style = {
        'font': 'Gentium',
        'size': 12
    }

    document = Document()
    styles = document.styles

    # TIBETAN
    bo_style = styles.add_style('Tibetan', WD_STYLE_TYPE.CHARACTER)
    bo_font = bo_style.font
    bo_font.name = tibetan_style['font']
    bo_font.size = Pt(tibetan_style['size'])
    # PEYDURMA NOTES
    prefix_style = styles.add_style('Peydurma Notes', WD_STYLE_TYPE.CHARACTER)
    prefix_font = prefix_style.font
    c = pedurma_style['color']
    prefix_font.color.rgb = RGBColor(c[0], c[1], c[2])

    # COMMUNICATIVE VERSION
    com_style = styles.add_style('Communicative', WD_STYLE_TYPE.CHARACTER)
    com_font = com_style.font
    com_font.name = communicative_style['font']
    com_font.size = Pt(communicative_style['size'])

    # SEMANTIC VERSION
    sem_style = styles.add_style('Semantic', WD_STYLE_TYPE.CHARACTER)
    sem_style.base_style = styles['Normal']
    sem_font = sem_style.font
    sem_font.name = semantic_style['font']
    sem_font.size = Pt(semantic_style['size'])

    # COMMUNICATIVE PARAGRAPH
    com_par_style = styles.add_style('Com. paragraph', WD_STYLE_TYPE.PARAGRAPH)
    com_par_style.paragraph_format.space_before = Cm(0)
    com_par_style.paragraph_format.space_after = Cm(0)

    # OTHER PARAGRAPH
    other_par_style = styles.add_style('Other paragraph', WD_STYLE_TYPE.PARAGRAPH)
    other_par_style.paragraph_format.space_before = Cm(0)
    other_par_style.paragraph_format.space_after = Cm(1)
    other_par_style.paragraph_format.left_indent = Cm(indentation)
    other_par_style.paragraph_format.line_spacing = WD_LINE_SPACING.SINGLE

    return document, word1, word2


def create_docx(lines, path):
    document, word1, word2 = initiate_docx()

    com_p = document.add_paragraph()
    for line in lines:
        chunks = get_syllables(line)
        for switch, chunk in chunks:
            for code, part in chunk:
                run = com_p.add_run(part)
                if switch == 1:
                    c = word1[code]['color']
                if switch == 2:
                    c = word2[code]['color']
                run.font.color.rgb = RGBColor(c[0], c[1], c[2])
        com_p.add_run().add_break()

    out_path = path.parent / (path.stem + '.docx')
    document.save(str(out_path))


in_file = Path('tib/soldep.txt')
lines = in_file.read_text(encoding='utf-8').strip().split('\n')
output = Path('out') / (in_file.stem + '.docx')
create_docx(lines, output)
